import streamlit as st
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
import pypdf
from docx import Document
from fpdf import FPDF
from io import BytesIO
import requests
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi
import yt_dlp
import os
import time
import datetime
from langchain_community.tools import DuckDuckGoSearchRun

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="StratIntel (Master)", page_icon="♟️", layout="wide")

# ==========================================
# 🔐 SISTEMA DE LOGIN
# ==========================================
def check_password():
    def password_entered():
        if st.session_state["username"] in st.secrets["passwords"] and \
           st.session_state["password"] == st.secrets["passwords"][st.session_state["username"]]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if st.session_state.get("password_correct", False):
        return True

    st.markdown("## ♟️ StratIntel: Acceso Restringido")
    st.text_input("Usuario", key="username")
    st.text_input("Contraseña", type="password", on_change=password_entered, key="password")
    
    if "password_correct" in st.session_state and not st.session_state["password_correct"]:
        st.error("❌ Credenciales inválidas")
    return False

if not check_password():
    st.stop()

# ==========================================
# ⚙️ CONFIGURACIÓN Y MODELO
# ==========================================
API_KEY_FIJA = "" 
if "GOOGLE_API_KEY" in st.secrets:
    API_KEY_FIJA = st.secrets["GOOGLE_API_KEY"]

MODELO_ACTUAL = "gemini-2.5-flash"  

# ==========================================
# 🧠 BASE DE DATOS MAESTRA (V25 - TOTAL DOMAIN AWARENESS)
# ==========================================
DB_CONOCIMIENTO = {
    "✨ RECOMENDACIÓN AUTOMÁTICA": {
        "desc": "La IA decide la mejor estrategia basándose en el contenido.",
        "preguntas": ["Identifica los hallazgos estratégicos más críticos.", "Realiza una evaluación integral de riesgos.", "Genera un Resumen Ejecutivo (BLUF).", "¿Cuáles son las anomalías o patrones ocultos más relevantes?"]
    },

    # =========================================================================
    # 🌍 BLOQUE 1: REALISMO, GEOPOLÍTICA Y PODER (HARD POWER)
    # =========================================================================
    "--- REALISMO, PODER Y ESPACIO ---": { "desc": "", "preguntas": [] },

    "Hans Morgenthau (Realismo Clásico Integral)": {
        "desc": "Los 6 Principios del Realismo Político y el Interés como Poder.",
        "preguntas": [
            "Leyes Objetivas: ¿Qué fuerzas inherentes a la naturaleza humana (egoísmo, dominio) están impulsando este conflicto?",
            "Interés y Poder: Define el 'Interés Nacional' de los actores en términos de poder, no de moralidad.",
            "Supervivencia del Estado: ¿Está la integridad territorial o política del Estado en riesgo directo?",
            "Autonomía de la Esfera Política: Analiza la decisión desde una lógica puramente política, ignorando consideraciones económicas o legales secundarias."
        ]
    },
    "Kenneth Waltz (Neorrealismo Estructural)": {
        "desc": "Las Tres Imágenes y la estructura anárquica del sistema.",
        "preguntas": [
            "Tercera Imagen (Sistémica): ¿Cómo la anarquía internacional y la polaridad obligan al actor a actuar así?",
            "Distribución de Capacidades: ¿El sistema es unipolar, bipolar o multipolar y cómo afecta esto la estabilidad?",
            "Equilibrio de Poder: ¿Está el actor haciendo 'Balancing' (aliarse contra el fuerte) o 'Bandwagoning' (unirse al fuerte)?",
            "Principio de Autoayuda: ¿Qué medidas unilaterales está tomando el actor para garantizar su seguridad?"
        ]
    },
    "John Mearsheimer (Realismo Ofensivo)": {
        "desc": "Hegemonía, Tragedia de las Potencias y Geografía.",
        "preguntas": [
            "Búsqueda de Hegemonía: ¿Está el actor intentando convertirse en el Hegemon regional para asegurar su supervivencia?",
            "Poder Detenedor del Agua: ¿Cómo la geografía (océanos, montañas) limita o facilita la proyección de poder?",
            "Maximizador de Poder: ¿Está el actor aprovechando oportunidades para debilitar a rivales potenciales?",
            "Estrategia de 'Buck-Passing': ¿Está intentando que otro estado asuma el costo de contener al agresor?"
        ]
    },
    "Halford Mackinder (Teoría del Heartland)": {
        "desc": "El control de la Isla Mundial y el Pivote Geográfico.",
        "preguntas": [
            "Pivote Geográfico: ¿Quién controla actualmente el 'Heartland' (Eurasia central)?",
            "Cinturón Interior: ¿Hay un conflicto por el control de las zonas costeras que rodean el Heartland?",
            "Contención Terrestre: ¿Se está usando el poder terrestre para negar el acceso a las potencias marítimas?"
        ]
    },
    "Nicholas Spykman (Teoría del Rimland)": {
        "desc": "El control de los bordes anfibios (Rimland).",
        "preguntas": [
            "Anfibia Estratégica: Analiza el conflicto en las zonas costeras/peninsulares (Rimland).",
            "Cerco: ¿Están las potencias tratando de rodear al actor central desde el mar?",
            "Valor de las Alianzas: ¿Qué alianzas en el borde euroasiático son vitales para mantener el equilibrio?"
        ]
    },
    "Realismo Defensivo (Walt & Jervis)": {
        "desc": "Equilibrio de Amenazas y Dilema de Seguridad.",
        "preguntas": [
            "Equilibrio de Amenazas: Evalúa la amenaza combinando: Poder Agregado, Geografía, Capacidad Ofensiva e Intenciones.",
            "Dilema de Seguridad: ¿Las medidas defensivas de un actor están siendo malinterpretadas como ofensivas?",
            "Espiral de Conflicto: ¿Cómo una acción defensiva ha provocado una reacción hostil involuntaria?"
        ]
    },
    "Realismo Periférico (Escudé) & Neoclásico (Schweller)": {
        "desc": "Estrategias para estados dependientes y filtros domésticos.",
        "preguntas": [
            "Costo-Beneficio de la Soberanía: ¿El costo de confrontar al Hegemon supera los beneficios ciudadanos?",
            "Política de Alineamiento: ¿Debería el estado adoptar un perfil bajo para obtener recursos?",
            "Coherencia Estatal: ¿Están las élites fragmentadas impidiendo una respuesta unificada al sistema internacional?"
        ]
    },

    # =========================================================================
    # 🤝 BLOQUE 2: LIBERALISMO, CONSTRUCTIVISMO Y SOFT POWER
    # =========================================================================
    "--- INSTITUCIONES, IDENTIDAD Y COOPERACIÓN ---": { "desc": "", "preguntas": [] },

    "Joseph Nye (Poder Multidimensional & Smart Power)": {
        "desc": "Interacción entre Hard, Soft y Smart Power.",
        "preguntas": [
            "Soft Power: ¿Qué activos de cultura, valores o políticas otorgan atracción y legitimidad?",
            "Smart Power: ¿Se combina eficazmente la coerción militar/económica con la persuasión diplomática?",
            "Tablero Tridimensional: Analiza el poder en sus tres niveles: Militar (Unipolar), Económico (Multipolar) y Transnacional (Difuso)."
        ]
    },
    "Keohane & Nye (Interdependencia Compleja)": {
        "desc": "Vínculos transnacionales y costo de ruptura.",
        "preguntas": [
            "Canales Múltiples: ¿Existen conexiones entre sociedades (no solo gobiernos) que afecten el conflicto?",
            "Sensibilidad y Vulnerabilidad: ¿Cuán costosa sería la ruptura de las relaciones económicas?",
            "Regímenes Internacionales: ¿Qué normas o reglas implícitas están gobernando las expectativas?"
        ]
    },
    "Robert Axelrod (Evolución de la Cooperación)": {
        "desc": "Teoría de Juegos aplicada a la cooperación sostenida.",
        "preguntas": [
            "Sombra del Futuro: ¿Es la interacción lo suficientemente duradera para fomentar la cooperación?",
            "Tit-for-Tat: ¿Está el actor respondiendo con reciprocidad estricta (cooperar si cooperas, castigar si traicionas)?",
            "Detección de Trampas: ¿Qué mecanismos de verificación existen para asegurar el cumplimiento?"
        ]
    },
    "Alexander Wendt (Constructivismo Social)": {
        "desc": "Identidades, normas y culturas de la anarquía.",
        "preguntas": [
            "Estructura Ideacional: ¿Cómo las identidades históricas definen el interés nacional?",
            "Culturas de la Anarquía: ¿El sistema es Hobbesiano (Enemigos), Lockeano (Rivales) o Kantiano (Amigos)?",
            "Ciclo de Refuerzo: ¿Cómo las interacciones pasadas han construido la percepción actual de 'amenaza'?"
        ]
    },
    "Samuel Huntington (Choque de Civilizaciones)": {
        "desc": "Conflictos de identidad cultural profunda.",
        "preguntas": [
            "Líneas de Falla: ¿Ocurre el conflicto en la frontera entre dos civilizaciones distintas?",
            "Síndrome del País Pariente: ¿Están otros estados interviniendo por lealtad cultural/religiosa?",
            "Occidente vs El Resto: ¿Es una reacción contra la imposición de valores occidentales?"
        ]
    },

    # =========================================================================
    # ⚔️ BLOQUE 3: ESTRATEGIA MILITAR, HÍBRIDA Y ASIMÉTRICA
    # =========================================================================
    "--- ARTE DE LA GUERRA Y CONFLICTO ---": { "desc": "", "preguntas": [] },

    "Sun Tzu (El Arte de la Guerra)": {
        "desc": "Engaño, velocidad y victoria sin combate.",
        "preguntas": [
            "El Engaño: ¿Toda la operación se basa en una finta o distracción?",
            "Ganar sin luchar: ¿Está el actor logrando sus objetivos políticos sin uso cinético de fuerza?",
            "Conocimiento: ¿Conoce el actor al enemigo y a sí mismo? (Evaluar brechas de inteligencia).",
            "Terreno: ¿Es el terreno mortal, disperso o clave? ¿Cómo afecta la maniobra?"
        ]
    },
    "Carl von Clausewitz (La Guerra Absoluta)": {
        "desc": "La trinidad y el Centro de Gravedad.",
        "preguntas": [
            "Trinidad Paradójica: Analiza el equilibrio entre Pasión (Pueblo), Probabilidad (Ejército) y Razón (Gobierno).",
            "Niebla y Fricción: ¿Qué imprevistos y falta de información están ralentizando la operación?",
            "Centro de Gravedad (COG): ¿Cuál es la fuente de poder del enemigo que, si cae, todo el sistema colapsa?"
        ]
    },
    "Guerra Híbrida (Doctrina Gerasimov)": {
        "desc": "Sincronización de medios militares y no militares.",
        "preguntas": [
            "Fase Latente: ¿Se está usando desinformación para desestabilizar la población antes del conflicto abierto?",
            "Fuerzas Proxy: ¿Se utilizan actores no estatales o mercenarios para negar responsabilidad (Plausible Deniability)?",
            "Dominio de la Información: ¿Es el ataque informativo más devastador que el ataque físico?"
        ]
    },
    "Qiao Liang & Wang Xiangsui (Guerra Irrestricta)": {
        "desc": "Todo es un arma: leyes, economía, drogas.",
        "preguntas": [
            "Desbordamiento del Campo de Batalla: ¿Se está usando el Lawfare (leyes) o la demografía como arma?",
            "Guerra Financiera: ¿Se están atacando las monedas o mercados del adversario?",
            "Guerra Cultural: ¿Se están atacando los valores fundacionales de la sociedad objetivo?"
        ]
    },

    # =========================================================================
    # 🧠 BLOQUE 4: PSICOLOGÍA OPERATIVA Y TOMA DE DECISIONES
    # =========================================================================
    "--- MENTE, LIDERAZGO Y DECISIÓN ---": { "desc": "", "preguntas": [] },

    "Graham Allison (Modelos de Decisión)": {
        "desc": "Lentes para analizar crisis gubernamentales.",
        "preguntas": [
            "Modelo I (Actor Racional): ¿Cuál es la opción lógica que maximiza beneficios estratégicos?",
            "Modelo II (Proceso Organizacional): ¿Qué rutinas y procedimientos estándar (SOPs) limitan la flexibilidad?",
            "Modelo III (Política Burocrática): ¿Qué luchas de poder internas entre agencias definieron la decisión?"
        ]
    },
    "Perfilado Dark Triad (Tríada Oscura)": {
        "desc": "Psicopatía, Narcisismo y Maquiavelismo en el liderazgo.",
        "preguntas": [
            "Narcisismo: ¿El líder necesita admiración constante y reacciona con ira desproporcionada a la crítica?",
            "Maquiavelismo: ¿El líder manipula a aliados y enemigos sin remordimiento para fines personales?",
            "Psicopatía: ¿Muestra falta total de empatía y toma riesgos impulsivos peligrosos?",
            "Vulnerabilidad del Ego: ¿Cómo se puede explotar su necesidad de validación?"
        ]
    },
    "Código MICE (Motivaciones de Traición)": {
        "desc": "Money, Ideology, Coercion, Ego.",
        "preguntas": [
            "Dinero (Money): ¿Existen crisis financieras personales o avaricia desmedida?",
            "Ideología (Ideology): ¿Cree el sujeto en una causa superior opuesta a su sistema actual?",
            "Coerción (Coercion): ¿Existe material comprometedor (Kompromat) para chantaje?",
            "Ego: ¿Se siente infravalorado o busca venganza contra sus superiores?"
        ]
    },
    "Gustave Le Bon (Psicología de Masas)": {
        "desc": "Comportamiento irracional y contagio emocional.",
        "preguntas": [
            "Contagio Mental: ¿Cómo se está propagando la emoción irracional en la población?",
            "Líder de Masas: ¿Quién está canalizando el odio o la esperanza de la multitud?",
            "Imágenes Simplistas: ¿Qué eslóganes o símbolos están reemplazando el pensamiento lógico?"
        ]
    },
    "Barry Buzan (Securitización)": {
        "desc": "La construcción de amenazas existenciales.",
        "preguntas": [
            "Actor Securitizador: ¿Quién declara el asunto como una 'amenaza existencial'?",
            "Objeto Referente: ¿Qué se intenta proteger (Estado, Identidad, Economía)?",
            "Medidas Extraordinarias: ¿Se usa la retórica de seguridad para justificar acciones fuera de la ley?"
        ]
    },
    "John Boyd (Ciclo OODA)": {
        "desc": "Velocidad de procesamiento: Observar, Orientar, Decidir, Actuar.",
        "preguntas": [
            "Velocidad del Ciclo: ¿Quién completa su ciclo de decisión más rápido?",
            "Fase de Orientación: ¿Cómo los sesgos culturales moldean la percepción de la realidad?",
            "Colapso: ¿Cómo generar ambigüedad para aislar al enemigo de su entorno?"
        ]
    },

    # =========================================================================
    # 💰 BLOQUE 5: GEOECONOMÍA Y TEORÍA DE JUEGOS
    # =========================================================================
    "--- GEOECONOMÍA Y RECURSOS ---": { "desc": "", "preguntas": [] },

    "Edward Luttwak (Geoeconomía)": {
        "desc": "La lógica del conflicto con la gramática del comercio.",
        "preguntas": [
            "Armamentalización del Comercio: ¿Se usan aranceles, sanciones o bloqueos como armas de guerra?",
            "Predación de Inversiones: ¿Está un estado adquiriendo infraestructura crítica del rival?",
            "Soberanía Tecnológica: ¿Se está bloqueando el acceso a chips, IA o tecnología clave?"
        ]
    },
    "Teoría de Juegos (John Nash & Schelling)": {
        "desc": "Equilibrios, Suma Cero y Disuasión.",
        "preguntas": [
            "Suma Cero vs Suma Variable: ¿Para que uno gane, el otro debe perderlo todo?",
            "Equilibrio de Nash: ¿Cuál es la situación donde nadie tiene incentivos para cambiar su estrategia?",
            "Juego de la Gallina (Chicken): ¿Quién cederá primero ante la inminencia del choque catastrófico?",
            "Credibilidad de la Amenaza: ¿Es creíble la promesa de castigo del actor?"
        ]
    },

    # =========================================================================
    # 🌐 BLOQUE 6: CIBERESPACIO Y REDES TECNOLÓGICAS
    # =========================================================================
    "--- CIBERINTELIGENCIA ---": { "desc": "", "preguntas": [] },

    "Cyber Kill Chain (Lockheed Martin)": {
        "desc": "Fases secuenciales de una intrusión.",
        "preguntas": [
            "Reconocimiento: ¿Qué datos se están recolectando antes del ataque?",
            "Armamentización y Entrega: ¿Cómo se creó y entregó el malware (Phishing, USB, Exploit)?",
            "Explotación e Instalación: ¿Qué vulnerabilidad técnica o humana se aprovechó?",
            "Acciones sobre Objetivos: ¿Se busca robo de datos (Espionaje), destrucción (Wiper) o secuestro (Ransomware)?"
        ]
    },
    "Teoría del Actor-Red (Latour)": {
        "desc": "Agencia de los objetos y algoritmos.",
        "preguntas": [
            "Agencia Tecnológica: ¿Cómo un algoritmo o plataforma está moldeando el conflicto por sí solo?",
            "Cajas Negras: ¿Qué procesos técnicos se están aceptando sin cuestionar su funcionamiento?",
            "Traducción de Intereses: ¿Cómo se redefinen los objetivos políticos al pasar por la red tecnológica?"
        ]
    },

    # =========================================================================
    # 🔮 BLOQUE 7: PROSPECTIVA, CAOS Y FUTUROS
    # =========================================================================
    "--- FUTUROS Y COMPLEJIDAD ---": { "desc": "", "preguntas": [] },

    "Análisis Causal por Capas (CLA - Inayatullah)": {
        "desc": "Deconstrucción profunda de la realidad.",
        "preguntas": [
            "La Letanía: ¿Qué dicen los titulares oficiales y datos superficiales?",
            "Causas Sistémicas: ¿Qué estructuras económicas, políticas o legales generan el problema?",
            "Visión del Mundo (Worldview): ¿Qué ideologías profundas sostienen el sistema actual?",
            "Mito y Metáfora: ¿Cuál es la historia inconsciente o arquetipo cultural que mueve a la sociedad?"
        ]
    },
    "Nassim Taleb (Cisne Negro & Antifragilidad)": {
        "desc": "Gestión de lo improbable y el caos.",
        "preguntas": [
            "Cisne Negro: Describe un evento de probabilidad casi nula pero impacto sistémico total.",
            "Rinoceronte Gris: ¿Qué amenaza obvia y visible estamos ignorando voluntariamente?",
            "Antifragilidad: ¿Qué actor se beneficia del desorden y se fortalece con el estrés?",
            "Falacia Narrativa: ¿Estamos inventando una historia coherente para datos que son puro ruido?"
        ]
    },
    "Análisis de Señales Débiles (Weak Signals)": {
        "desc": "Detección temprana de anomalías.",
        "preguntas": [
            "Ruido Marginal: ¿Qué dato 'irrelevante' se repite sospechosamente en contextos distintos?",
            "Filtro de Expertos: ¿Qué escenario están descartando los expertos por considerarlo 'imposible'?",
            "Patrones de Rareza: ¿Qué evento rompe la continuidad histórica establecida?"
        ]
    },

    # =========================================================================
    # 🛠️ BLOQUE 8: TÉCNICAS ESTRUCTURADAS DE ANÁLISIS (SATs)
    # =========================================================================
    "--- HERRAMIENTAS TÁCTICAS (SATs) ---": { "desc": "", "preguntas": [] },

    "Análisis de Hipótesis en Competencia (ACH)": {
        "desc": "Matriz científica para evitar sesgos.",
        "preguntas": [
            "Generación: Formula 4 hipótesis exclusivas.",
            "Evidencia: Lista toda la evidencia disponible.",
            "Diagnóstico: Evalúa la consistencia de cada evidencia con cada hipótesis.",
            "Refutación: Busca la evidencia que DESCARTE hipótesis, no que las confirme."
        ]
    },
    "Análisis de Actores (Stakeholder Mapping)": {
        "desc": "Mapa de poder e intereses.",
        "preguntas": [
            "Matriz Poder/Interés: Clasifica a los actores clave.",
            "Vetadores: ¿Quién tiene capacidad de bloqueo?",
            "Spoilers: ¿Quién se beneficia de que el conflicto continúe?"
        ]
    },
    "Matriz CARVER (Selección de Objetivos)": {
        "desc": "Evaluación de blancos para operaciones.",
        "preguntas": [
            "Criticidad: ¿Qué tan vital es para la misión?",
            "Accesibilidad: ¿Qué tan fácil es llegar al objetivo?",
            "Recuperabilidad: ¿Cuánto tiempo tardan en reemplazarlo?",
            "Vulnerabilidad: ¿Qué recursos se necesitan para dañarlo?",
            "Efecto: ¿Cuál es el impacto sistémico?",
            "Reconocibilidad: ¿Se puede identificar fácilmente?"
        ]
    },
    "Análisis PMESII-PT (Entorno Operativo)": {
        "desc": "Análisis holístico del teatro de operaciones.",
        "preguntas": [
            "Político y Militar.",
            "Económico y Social.",
            "Información e Infraestructura.",
            "Entorno Físico y Tiempo."
        ]
    },
    "Análisis DIME (Instrumentos de Poder)": {
        "desc": "Capacidades nacionales.",
        "preguntas": [
            "Diplomático: Alianzas y aislamiento.",
            "Informacional: Narrativa y ciber.",
            "Militar: Disuasión y fuerza.",
            "Económico: Sanciones y ayuda."
        ]
    },
    "Análisis FODA (SWOT) de Inteligencia": {
        "desc": "Ofensivo/Defensivo.",
        "preguntas": [
            "Amenazas Externas Inminentes.",
            "Oportunidades de Explotación.",
            "Debilidades Internas (Vulnerabilidades).",
            "Fortalezas (Capacidades)."
        ]
    },
    "Técnica de los 5 Porqués": {
        "desc": "Búsqueda de Causa Raíz.",
        "preguntas": [
            "Sintoma visible.",
            "¿Por qué ocurre? (Nivel 1)",
            "¿Por qué ocurre el nivel anterior? (Repetir hasta Nivel 5)",
            "Falla Sistémica Raíz."
        ]
    },
    "Abogado del Diablo": {
        "desc": "Desafío de asunciones.",
        "preguntas": [
            "Desafío Frontal: Argumenta por qué la conclusión principal está equivocada.",
            "Defensa de lo Irracional: Asume que el adversario actuará de forma ilógica y explícalo."
        ]
    }
}

# --- GESTIÓN DE ESTADO ---
if 'api_key' not in st.session_state: st.session_state['api_key'] = ""
if 'texto_analisis' not in st.session_state: st.session_state['texto_analisis'] = ""
if 'origen_dato' not in st.session_state: st.session_state['origen_dato'] = "Ninguno"

# --- FUNCIONES DE PROCESAMIENTO ---
def buscar_en_web(query):
    try:
        search = DuckDuckGoSearchRun()
        return search.run(query)
    except Exception as e: return f"Error web: {e}"

def procesar_archivos_pdf(archivos):
    texto_total = ""
    nombres = []
    for archivo in archivos:
        reader = pypdf.PdfReader(archivo)
        texto_pdf = "".join([p.extract_text() for p in reader.pages])
        texto_total += f"\n--- ARCHIVO: {archivo.name} ---\n{texto_pdf}\n"
        nombres.append(archivo.name)
    return texto_total, str(nombres)

def procesar_archivos_docx(archivos):
    texto_total = ""
    nombres = []
    for archivo in archivos:
        doc = Document(archivo)
        texto_doc = "\n".join([para.text for para in doc.paragraphs])
        texto_total += f"\n--- ARCHIVO: {archivo.name} ---\n{texto_doc}\n"
        nombres.append(archivo.name)
    return texto_total, str(nombres)

def obtener_texto_web(url):
    try:
        h = {'User-Agent': 'Mozilla/5.0'}
        r = requests.get(url, headers=h, timeout=15)
        s = BeautifulSoup(r.content, 'html.parser')
        for script in s(["script", "style"]): script.extract()
        return s.get_text(separator='\n')
    except Exception as e: return f"Error: {e}"

def procesar_youtube(url, api_key):
    vid = url.split("v=")[-1].split("&")[0] if "v=" in url else url.split("/")[-1]
    try:
        t = YouTubeTranscriptApi.get_transcript(vid, languages=['es', 'en'])
        return " ".join([i['text'] for i in t]), "Subtítulos"
    except:
        st.info(f"Multimodal (Audio)...")
        opts = {'format': 'bestaudio/best', 'outtmpl': '%(id)s.%(ext)s', 'postprocessors': [{'key': 'FFmpegExtractAudio','preferredcodec': 'mp3'}], 'quiet': True}
        try:
            with yt_dlp.YoutubeDL(opts) as ydl:
                info = ydl.extract_info(url, download=True)
                fname = f"{info['id']}.mp3"
            genai.configure(api_key=api_key)
            myfile = genai.upload_file(fname)
            while myfile.state.name == "PROCESSING": time.sleep(2); myfile = genai.get_file(myfile.name)
            model = genai.GenerativeModel(MODELO_ACTUAL)
            res = model.generate_content([myfile, "Transcribe el audio."])
            if os.path.exists(fname): os.remove(fname)
            myfile.delete()
            return res.text, "Audio IA"
        except Exception as e: return f"Error: {e}", "Error"

# --- FUNCIONES DE REPORTE ---
def limpiar_texto(t):
    if not t: return ""
    reps = {"✨": "", "🚀": "", "⚠️": "[!]", "✅": "[OK]", "🛡️": "", "🔒": "", "🎖️": "", "♟️": "", "⚖️": ""}
    for k,v in reps.items(): t = t.replace(k,v)
    return t.encode('latin-1', 'replace').decode('latin-1')

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'StratIntel Report V16', 0, 1, 'C')
        self.ln(5)
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 7)
        self.cell(0, 10, 'Generado por IA. Uso Confidencial.', 0, 0, 'C')

def crear_pdf(texto, tecnicas, fuente):
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_font("Arial", "B", 10)
    pdf.multi_cell(0, 5, limpiar_texto(f"Fuente: {fuente}\nTécnicas: {tecnicas}"))
    pdf.ln(5)
    pdf.set_font("Arial", "", 10)
    pdf.multi_cell(0, 5, limpiar_texto(texto))
    return pdf.output(dest='S').encode('latin-1', 'replace')

def crear_word(texto, tecnicas, fuente):
    doc = Document()
    doc.add_heading('StratIntel Intelligence Report', 0)
    doc.add_paragraph(f"Fuente: {fuente}").bold = True
    doc.add_paragraph(f"Técnicas: {tecnicas}").bold = True
    for l in texto.split('\n'):
        if l.startswith('#'): doc.add_heading(l.replace('#','').strip(), level=2)
        else: doc.add_paragraph(l)
    
    aviso = doc.add_paragraph()
    aviso.add_run("\n\n------------------\nAVISO: Generado por IA. Verificar datos.").font.size = 8
    b = BytesIO(); doc.save(b); b.seek(0)
    return b

# --- INTERFAZ ---
st.sidebar.title("♟️ StratIntel")
st.sidebar.caption("Master Edition | Ops Mode")
st.sidebar.markdown("---")

if API_KEY_FIJA:
    st.session_state['api_key'] = API_KEY_FIJA
    genai.configure(api_key=API_KEY_FIJA)
    st.sidebar.success(f"✅ Conectado ({MODELO_ACTUAL})")
else:
    if not st.session_state['api_key']:
        k = st.sidebar.text_input("🔑 API KEY:", type="password")
        if k: st.session_state['api_key'] = k; genai.configure(api_key=k); st.rerun()

# SELECTOR MULTI-TECNICA
st.sidebar.subheader("🎯 Misión")
tecnicas_seleccionadas = st.sidebar.multiselect(
    "Técnicas (Máx 3):",
    options=list(DB_CONOCIMIENTO.keys()),
    max_selections=3
)

temp = st.sidebar.slider("Creatividad", 0.0, 1.0, 0.4)
if st.sidebar.button("🔒 Salir"): del st.session_state["password_correct"]; st.rerun()

st.title("♟️ StratIntel | División de Análisis")
st.markdown("**Sistema de Inteligencia Estratégica (DSS)**")

# CARGA
t1, t2, t3, t4, t5 = st.tabs(["📂 PDFs", "📝 DOCXs", "🌐 Web", "📺 YouTube", "✍️ Manual"])
with t1:
    f = st.file_uploader("PDFs", type="pdf", accept_multiple_files=True)
    if f and st.button("Procesar PDF"):
        t, n = procesar_archivos_pdf(f); st.session_state['texto_analisis']=t; st.session_state['origen_dato']=f"PDFs: {n}"; st.success(f"✅ {len(f)}")
with t2:
    f = st.file_uploader("DOCXs", type="docx", accept_multiple_files=True)
    if f and st.button("Procesar DOCX"):
        t, n = procesar_archivos_docx(f); st.session_state['texto_analisis']=t; st.session_state['origen_dato']=f"DOCXs: {n}"; st.success(f"✅ {len(f)}")
with t3:
    u = st.text_input("URL"); 
    if st.button("Web"): st.session_state['texto_analisis']=obtener_texto_web(u); st.session_state['origen_dato']=f"Web: {u}"; st.success("OK")
with t4:
    y = st.text_input("YouTube")
    if st.button("Video"):
        with st.spinner("..."):
            t,m=procesar_youtube(y,st.session_state['api_key'])
            if m!="Error": st.session_state['texto_analisis']=t; st.session_state['origen_dato']=f"YT: {y}"; st.success("OK")
            else: st.error(t)
with t5:
    m = st.text_area("Manual")
    if st.button("Fijar"): st.session_state['texto_analisis']=m; st.session_state['origen_dato']="Manual"; st.success("OK")

st.markdown("---")
if st.session_state['texto_analisis']:
    with st.expander(f"Fuente Activa: {st.session_state['origen_dato']}"): st.write(st.session_state['texto_analisis'][:1000])

# EJECUCIÓN
st.header("Generación de Informe")

if not st.session_state['api_key'] or not st.session_state['texto_analisis']:
    st.warning("⚠️ Carga datos para comenzar.")
else:
    c1, c2 = st.columns([1, 2])
    with c1:
        if not tecnicas_seleccionadas: st.info("👈 Selecciona técnicas.")
        
        # --- SELECTOR DE PROFUNDIDAD CON MODO OPERACIONAL ---
        profundidad = st.radio(
            "Nivel de Profundidad:", 
            ["🔍 Estratégico (Resumen)", "🎯 Táctico (Todas las preguntas)", "⚙️ Operacional (Selección Específica)"],
            help="Estratégico: Visión general. Táctico: Todas las preguntas del marco. Operacional: Selecciona preguntas manualmente."
        )
        
        # --- LÓGICA DE SELECCIÓN MANUAL (OPERACIONAL) ---
        preguntas_manuales = {}
        if "Operacional" in profundidad and tecnicas_seleccionadas:
            st.info("👇 Selecciona los vectores de análisis:")
            for tec in tecnicas_seleccionadas:
                # Obtenemos las preguntas de TU base de datos exacta
                qs = DB_CONOCIMIENTO.get(tec, {}).get("preguntas", [])
                if qs:
                    sel = st.multiselect(f"Preguntas para {tec}:", qs)
                    preguntas_manuales[tec] = sel
                else:
                    st.warning(f"{tec} no tiene preguntas predefinidas.")
        
        usar_internet = st.checkbox("🌐 Búsqueda Web")
        pir = st.text_area("PIR (Opcional):", height=100)

    with c2:
        if st.button("🚀 EJECUTAR MISIÓN", type="primary", use_container_width=True, disabled=len(tecnicas_seleccionadas)==0):
            try:
                genai.configure(api_key=st.session_state['api_key'])
                model = genai.GenerativeModel(MODELO_ACTUAL)
                ctx = st.session_state['texto_analisis']
                
                # BÚSQUEDA WEB
                contexto_web = ""
                if usar_internet:
                    with st.status("🌐 Buscando...", expanded=True) as s:
                        q = f"{pir} {st.session_state['origen_dato']}" if pir else f"Análisis {st.session_state['origen_dato']}"
                        res_web = buscar_en_web(q)
                        contexto_web = f"\nINFO WEB:\n{res_web}\n"
                        s.update(label="✅ Hecho", state="complete", expanded=False)
                
                # BUCLE DE ANÁLISIS
                informe_final = f"# INFORME\nFECHA: {datetime.datetime.now().strftime('%d/%m/%Y')}\nFUENTE: {st.session_state['origen_dato']}\n\n"
                progreso = st.progress(0)
                
                for i, tec in enumerate(tecnicas_seleccionadas):
                    st.caption(f"Analizando: {tec}...")
                    
                    # LÓGICA DE INYECCIÓN DE PREGUNTAS
                    instruccion_preguntas = ""
                    
                    if "Táctico" in profundidad:
                        qs = DB_CONOCIMIENTO.get(tec, {}).get("preguntas", [])
                        if qs:
                            lista = "\n".join([f"- {p}" for p in qs])
                            instruccion_preguntas = f"\n\nOBLIGATORIO: Responde DETALLADAMENTE a TODAS estas preguntas del marco teórico:\n{lista}"
                        else:
                            instruccion_preguntas = "\n\nINSTRUCCIÓN: Realiza un análisis táctico detallado."

                    elif "Operacional" in profundidad:
                        qs_selec = preguntas_manuales.get(tec, [])
                        if qs_selec:
                            lista = "\n".join([f"- {p}" for p in qs_selec])
                            instruccion_preguntas = f"\n\nOBLIGATORIO: Centra el análisis EXCLUSIVAMENTE en responder estas preguntas seleccionadas:\n{lista}"
                        else:
                            instruccion_preguntas = "\n\n(NOTA: El usuario no seleccionó preguntas específicas. Realiza un análisis general de la técnica)."

                    else: # Estratégico
                        instruccion_preguntas = "\n\nINSTRUCCIÓN: Realiza un análisis estratégico general, fluido y ejecutivo (Resumen Global)."

                    prompt = f"""
                    ACTÚA COMO: Analista de Inteligencia Senior y Experto en Relaciones Internacionales.
                    METODOLOGÍA: {tec}
                    PIR (Requerimiento de Inteligencia): {pir}
                    
                    {instruccion_preguntas}
                    
                    CONTEXTO DOCUMENTAL:
                    {ctx}
                    {contexto_web}
                    
                    FORMATO: Académico, riguroso, citar fuentes del texto.
                    """
                    
                    # RETRY LOGIC
                    intentos = 0
                    exito = False
                    while intentos < 3 and not exito:
                        try:
                            res = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(temperature=temp))
                            informe_final += f"\n\n## 📌 {tec}\n{res.text}\n\n---\n"
                            exito = True
                        except Exception as e:
                            if "429" in str(e):
                                st.warning(f"⚠️ Tráfico alto (429). Esperando 30s... (Intento {intentos+1})")
                                time.sleep(30)
                                intentos += 1
                            else:
                                st.error(f"Error: {e}")
                                break

                    progreso.progress((i + 1) / len(tecnicas_seleccionadas))
                    time.sleep(5) 
                
                st.session_state['res'] = informe_final
                st.session_state['tecnicas_usadas'] = ", ".join(tecnicas_seleccionadas)
                st.success("✅ Misión Completada")
                st.markdown(informe_final)

            except Exception as e: st.error(f"Error: {e}")

if 'res' in st.session_state:
    st.markdown("---")
    c1, c2 = st.columns(2)
    c1.download_button("Descargar Word", crear_word(st.session_state['res'], st.session_state['tecnicas_usadas'], st.session_state['origen_dato']), "Reporte.docx")
    try: c2.download_button("Descargar PDF", bytes(crear_pdf(st.session_state['res'], st.session_state['tecnicas_usadas'], st.session_state['origen_dato'])), "Reporte.pdf")
    except: pass

